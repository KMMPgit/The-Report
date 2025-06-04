import os
from docx import Document
import pdfkit
from jinja2 import Template
from datetime import datetime
import zipfile

def generate_documents(topic: str, summary: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_topic = topic.replace(" ", "_")
    base_filename = f"report_{safe_topic}_{timestamp}"
    docx_filename = f"{base_filename}.docx"
    pdf_filename = f"{base_filename}.pdf"
    zip_filename = f"{base_filename}.zip"

    doc = Document()
    doc.add_heading(f"Report: {topic}", 0)
    doc.add_paragraph(summary)
    doc.save(docx_filename)

    html_template = Template("""
    <html>
    <head><meta charset="utf-8"><title>{{ topic }}</title></head>
    <body>
        <h1>Report: {{ topic }}</h1>
        <p style="white-space: pre-wrap;">{{ summary }}</p>
    </body>
    </html>
    """)
    html_content = html_template.render(topic=topic, summary=summary)

    with open("temp_report.html", "w", encoding="utf-8") as f:
        f.write(html_content)

    pdfkit.from_file("temp_report.html", pdf_filename)
    os.remove("temp_report.html")

    with zipfile.ZipFile(zip_filename, "w") as zipf:
        zipf.write(docx_filename)
        zipf.write(pdf_filename)

    os.remove(docx_filename)
    os.remove(pdf_filename)

    return zip_filename
